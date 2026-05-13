"""
resume_extractor.py
File → raw text → LLM → JSON Resume schema dict.
Supports .tex, .docx, and .pdf input.
"""

import io
import re
import json
import logging

log = logging.getLogger("resume")

# ── Empty schema (used as default / merge target) ─────────────────────────────

EMPTY_RESUME = {
    "basics": {
        "name": "", "email": "", "phone": "",
        "location": "", "linkedin": "", "github": "", "portfolio": "",
    },
    "summary": "",
    "experience": [],
    "education": [],
    "skills": [],
    "projects": [],
    "certifications": [],
    "publications": [],
    "awards": [],
    "volunteer": [],
    "languages": [],
}

# ── LLM system prompt ─────────────────────────────────────────────────────────

EXTRACTION_SYSTEM_PROMPT = """\
You are a professional resume parser. Extract ALL information from the resume \
text below and output it as a single JSON object following the exact schema provided.

RULES:
- Extract every piece of information present; do not omit anything
- Never invent or hallucinate information that is not in the resume
- If a field has no data, use empty string "" or empty list []
- Preserve original date formats (e.g. "June 2020", "Aug. 2018", "Present")
- Each bullet point becomes a separate string in the bullets array. Keep the \
full text of each bullet, do not truncate
- For skills: preserve the category groupings if they exist; otherwise create \
reasonable categories (e.g. "Languages", "Frameworks", "Tools", "Cloud")
- For linkedin/github/portfolio: include the URL or handle exactly as shown
- Return ONLY valid JSON with no commentary, no markdown, and no code fences

SCHEMA (fill every field):
{
  "basics": {
    "name": "full name",
    "email": "email address",
    "phone": "phone number",
    "location": "city, state or full address line",
    "linkedin": "linkedin URL or path e.g. linkedin.com/in/username",
    "github": "github URL or path e.g. github.com/username",
    "portfolio": "personal website or portfolio URL, empty string if none"
  },
  "summary": "summary or objective paragraph as a single string, empty if none",
  "experience": [
    {
      "company": "company or organization name",
      "title": "job title",
      "location": "city, state or country",
      "startDate": "start date string",
      "endDate": "end date string or Present",
      "bullets": ["full bullet text 1", "full bullet text 2"]
    }
  ],
  "education": [
    {
      "institution": "school or university name",
      "degree": "degree type e.g. Master of Science, Bachelor of Arts",
      "field": "field of study e.g. Computer Science",
      "startDate": "start date or empty string if not listed",
      "endDate": "graduation date",
      "gpa": "GPA if listed, empty string otherwise",
      "honors": ["honor or distinction if listed"],
      "location": "city, state or country"
    }
  ],
  "skills": [
    {
      "category": "category name e.g. Programming, Frameworks, Tools, Cloud, Databases",
      "items": ["skill1", "skill2", "skill3"]
    }
  ],
  "projects": [
    {
      "name": "project name",
      "startDate": "start date or empty string",
      "endDate": "end date or empty string",
      "bullets": ["description bullet 1", "description bullet 2"],
      "link": "project URL or repo link if listed, empty string otherwise"
    }
  ],
  "certifications": [
    {
      "name": "certification name",
      "issuer": "issuing organization",
      "date": "date obtained"
    }
  ],
  "publications": [
    {
      "title": "paper or article title",
      "journal": "journal or conference name",
      "date": "publication date",
      "link": "DOI or URL if listed"
    }
  ],
  "awards": ["award description as a single string"],
  "volunteer": ["volunteer role or activity description"],
  "languages": ["language name"]
}\
"""

# ── Text extractors ───────────────────────────────────────────────────────────

def extract_text_from_file(file_bytes: bytes, extension: str) -> str:
    """Route file bytes to the appropriate text extractor."""
    ext = extension.lower()
    if ext == ".tex":
        return _extract_tex(file_bytes)
    elif ext == ".docx":
        return _extract_docx(file_bytes)
    elif ext == ".pdf":
        return _extract_pdf(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext!r}. Accepted: .tex, .docx, .pdf")


def _extract_tex(file_bytes: bytes) -> str:
    """Strip LaTeX markup and return readable plain text."""
    text = file_bytes.decode("utf-8", errors="replace")

    # Keep only the document body
    if "\\begin{document}" in text:
        text = text.split("\\begin{document}", 1)[1]
    if "\\end{document}" in text:
        text = text.split("\\end{document}", 1)[0]

    # Preserve escaped special chars BEFORE comment stripping
    text = text.replace("\\%", "PCT")   # \% → literal percent sign
    text = text.replace("\\&", " and ")
    text = text.replace("\\$", "$")
    text = text.replace("\\_", "_")     # \_ → literal underscore (e.g. file_names)

    # Remove LaTeX line comments (% to end of line)
    text = re.sub(r"%.*$", "", text, flags=re.MULTILINE)

    # Restore preserved chars
    text = text.replace("PCT", "%")

    # Convert section headings → readable markers
    text = re.sub(r"\\(?:sub)*section\*?\{([^}]+)\}", r"\n\n=== \1 ===\n", text)

    # Custom two-arg environments: keep both args as readable text
    # e.g. \begin{joblong}{Title | Company, Loc}{Date}
    text = re.sub(
        r"\\begin\{[a-zA-Z]+\}\{([^}]*)\}\s*\{([^}]*)\}",
        lambda m: f"\n{m.group(1).strip()}  |  {m.group(2).strip()}\n",
        text,
    )
    # Single-arg environments: keep the arg
    text = re.sub(r"\\begin\{[a-zA-Z]+\}\{([^}]*)\}", r"\n\1\n", text)
    # Bare \begin{env} and \end{env} → remove
    text = re.sub(r"\\(?:begin|end)\{[a-zA-Z@*]+\}", "", text)

    # Convert \item → bullet dash
    text = re.sub(r"\\item\b\s*", "\n- ", text)

    # \\ with optional spacing arg [Xpt] → newline
    text = re.sub(r"\\\\\[?[0-9.a-z]*\]?", "\n", text)

    # Unwrap single-arg text formatting commands (keep the content)
    for cmd in [
        "textbf", "textit", "texttt", "emph", "small", "large", "Large",
        "huge", "Huge", "scshape", "normalfont", "mbox", "underline",
    ]:
        text = re.sub(r"\\" + cmd + r"\{([^}]*)\}", r"\1", text)

    # \href{url}{display} → display
    text = re.sub(r"\\href(?:WithoutArrow)?\{[^}]*\}\{([^}]*)\}", r"\1", text)

    # \hfill → tab-like gap (helps LLM see the right-aligned date)
    text = re.sub(r"\\hfill", "    ", text)

    # tabular column separator → space
    text = re.sub(r"\s*&\s*", "  ", text)

    # Strip all remaining LaTeX commands with optional [] and {} args
    text = re.sub(r"\\[a-zA-Z@]+\*?(?:\[[^\]]*\])?(?:\{[^}]*\})*", "", text)

    # Strip leftover braces, pipes, bare brackets like [2pt]
    text = re.sub(r"\[[0-9.]+[a-z]*\]", "", text)   # [2pt], [0.5in], etc.
    # Strip tabular column specs (e.g. "@{}l X@{}", "l X@") but not email @
    text = re.sub(r"@\{[^}]*\}", "", text)           # @{...} column padding
    text = re.sub(r"(?<!\w)l X(?!\w)", "", text)     # bare "l X" column spec
    text = re.sub(r"[{}|]", "", text)
    text = re.sub(r"\$[^$]*\$", "", text)

    # Clean up whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _extract_docx(file_bytes: bytes) -> str:
    """Extract all paragraph and table text from a DOCX file."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append("  ".join(cells))

    return "\n".join(lines)


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from all pages of a PDF."""
    import zlib
    from pypdf import PdfReader
    try:
        # Many real-world PDFs contain streams that can trigger zlib errors in strict mode.
        # strict=False makes parsing more tolerant in packaged builds.
        reader = PdfReader(io.BytesIO(file_bytes), strict=False)
    except Exception as e:
        raise ValueError(
            "We could not read this PDF. Please re-export it as a standard, text-based PDF (not a scanned image), then try again."
        ) from e

    pages: list[str] = []
    try:
        total = len(reader.pages)
    except Exception:
        total = 0

    if total > 0:
        for i in range(total):
            try:
                page = reader.pages[i]
                t = page.extract_text()
            except zlib.error:
                # zlib "incorrect header check" and similar: skip the page.
                continue
            except Exception:
                continue
            if t and t.strip():
                pages.append(t.strip())
    else:
        # Fallback iteration if len/pages indexing is not reliable for the PDF.
        try:
            for page in reader.pages:
                try:
                    t = page.extract_text()
                except zlib.error:
                    continue
                except Exception:
                    continue
                if t and t.strip():
                    pages.append(t.strip())
        except zlib.error as e:
            raise ValueError(
                "We could not read this PDF. Please re-export it as a standard, text-based PDF (not a scanned image), then try again."
            ) from e

    text = "\n\n".join(pages).strip()
    if not text:
        raise ValueError(
            "We could not extract text from this PDF. If it is a scanned PDF, please upload a text-based PDF or a DOCX file."
        )
    return text


# ── LLM call + JSON parsing ───────────────────────────────────────────────────

def _strip_fences(text: str) -> str:
    """Remove markdown code fences if the LLM wrapped its JSON in them."""
    text = text.strip()
    if text.startswith("```"):
        # Drop the opening fence line (e.g. ```json)
        text = text.split("\n", 1)[1] if "\n" in text else text[3:]
        if text.endswith("```"):
            text = text.rsplit("```", 1)[0]
    return text.strip()


def _fill_defaults(data: dict) -> dict:
    """Ensure all top-level schema keys exist with correct types."""
    result = {**EMPTY_RESUME, **data}
    if not isinstance(result.get("basics"), dict):
        result["basics"] = {}
    for key in EMPTY_RESUME["basics"]:
        result["basics"].setdefault(key, "")
    return result


def extract_resume_to_json(file_bytes: bytes, extension: str, llm_client) -> dict:
    """
    Full pipeline:
      file bytes  →  raw text  →  LLM  →  JSON Resume dict

    Args:
        file_bytes: raw bytes of the uploaded file
        extension:  file extension including dot, e.g. ".tex"
        llm_client: an LLMClient instance

    Returns:
        dict conforming to JSON Resume schema
    """
    # Step 1: extract readable text
    raw_text = extract_text_from_file(file_bytes, extension)
    log.info(f"[extract] text extracted from {extension} ({len(raw_text)} chars)")
    log.debug(f"[extract] text preview:\n{raw_text[:800]}")

    # Step 2: ask LLM to structure it
    user_prompt = f"Extract this resume into JSON:\n\n{raw_text}"
    response = llm_client.complete(
        EXTRACTION_SYSTEM_PROMPT,
        user_prompt,
        max_tokens=8192,
    )
    log.info("[extract] LLM responded")
    log.debug(f"[extract] raw LLM response:\n{response[:800]}")

    # Step 3: parse and fill defaults
    cleaned = _strip_fences(response)
    data = json.loads(cleaned)
    return _fill_defaults(data)
