import base64
import json
import logging
import os
import re
import subprocess
import tempfile
from pathlib import Path
from datetime import datetime

log = logging.getLogger("resume")

LIBREOFFICE_BIN = "/Applications/LibreOffice.app/Contents/MacOS/soffice"


def export_to_pdf(
    sections: dict,
    original_sections: dict | None,
    template_content: str | None,
    template_name: str | None,
    save_dir: str,
) -> str:
    ext = Path(template_name).suffix.lower() if template_name else ""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(save_dir, f"resume_{timestamp}.pdf")

    if template_content and ext == ".tex":
        return _export_tex_template(sections, original_sections, template_content, output_path)
    elif template_content and ext == ".docx":
        return _export_docx_template(sections, original_sections, template_content, output_path)
    else:
        return _export_plain_docx(sections, output_path)


def export_to_tex(
    sections: dict,
    original_sections: dict | None,
    template_content: str | None,
    save_dir: str,
) -> tuple[str, str]:
    """
    Save the original and edited .tex files side by side.
    Returns (original_path, edited_path).
    """
    if not template_content:
        raise ValueError("No .tex template content provided")

    b64 = template_content.split(",", 1)[1] if "," in template_content else template_content
    original_source = base64.b64decode(b64).decode("utf-8", errors="replace")

    edited_source = original_source
    if original_sections:
        orig_items = _flatten_sections(original_sections)
        edit_items = _flatten_sections(sections)
        for orig, edited in zip(orig_items, edit_items):
            if orig != edited:
                edited_source = _tex_replace(edited_source, orig, edited)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = Path(save_dir)
    out.mkdir(parents=True, exist_ok=True)

    orig_path = str(out / f"original_resume_{timestamp}.tex")
    edit_path = str(out / f"edited_resume_{timestamp}.tex")

    with open(orig_path, "w", encoding="utf-8") as f:
        f.write(original_source)
    with open(edit_path, "w", encoding="utf-8") as f:
        f.write(edited_source)

    log.info(f"TEX EXPORT — original: {orig_path}")
    log.info(f"TEX EXPORT — edited:   {edit_path}")

    return orig_path, edit_path


# ---------------------------------------------------------------------------
# LaTeX template export  — zero extra tokens, direct text substitution
# ---------------------------------------------------------------------------

def _export_tex_template(
    sections: dict,
    original_sections: dict | None,
    template_content: str,
    output_path: str,
) -> str:
    """
    Fill a LaTeX template by substituting edited text for original text in-place,
    then compile with pdflatex/xelatex. Uses zero LLM tokens.

    Strategy:
    1. Flatten original_sections and sections into parallel ordered lists of strings.
    2. For each (original, edited) pair that differs, find the original string in the
       .tex source and replace it with the LaTeX-escaped edited string.
    3. Normalization handles pandoc transformations (en-dashes, smart quotes, etc.)
       so that text extracted by pandoc can still be located in the raw .tex source.
    """
    b64 = template_content.split(",", 1)[1] if "," in template_content else template_content
    tex_source = base64.b64decode(b64).decode("utf-8", errors="replace")

    log.info("="*80)
    log.info("TEX EXPORT — RAW TEMPLATE AFTER DECODING")
    log.info("="*80)
    log.info(tex_source)
    log.info("="*80)

    if original_sections:
        orig_items = _flatten_sections(original_sections)
        edit_items = _flatten_sections(sections)

        log.info(f"TEX EXPORT — {len(orig_items)} original items, {len(edit_items)} edited items")
        log.info("-"*80)
        for i, (orig, edited) in enumerate(zip(orig_items, edit_items)):
            if orig == edited:
                log.info(f"[{i}] UNCHANGED: {repr(orig)}")
            else:
                found = orig in tex_source
                log.info(f"[{i}] ORIGINAL : {repr(orig)}")
                log.info(f"[{i}] NEW      : {repr(edited)}")
                log.info(f"[{i}] MATCH    : {'YES' if found else 'NO — NOT FOUND IN SOURCE'}")
                tex_source = _tex_replace(tex_source, orig, edited)
        log.info("-"*80)
    else:
        log.info("TEX EXPORT — no original_sections provided, skipping substitution")

    log.info("="*80)
    log.info("TEX EXPORT — FINAL .TEX CONTENT BEFORE COMPILATION")
    log.info("="*80)
    log.info(tex_source)
    log.info("="*80)

    out_dir = os.path.dirname(output_path)
    with tempfile.NamedTemporaryFile(
        suffix=".tex", delete=False, dir="/tmp", mode="w", encoding="utf-8"
    ) as tmp:
        tmp.write(tex_source)
        tmp_path = tmp.name

    try:
        pdf_path = _pdflatex_compile(tmp_path, out_dir)
        os.rename(pdf_path, output_path)
        return output_path
    finally:
        stem = Path(tmp_path).stem
        # .tex lives in /tmp; .aux/.log/.out are written to out_dir by pdflatex
        for f in [os.path.join("/tmp", stem + ".tex")]:
            try:
                os.unlink(f)
            except FileNotFoundError:
                pass
        for aux_ext in (".aux", ".log", ".out"):
            try:
                os.unlink(os.path.join(out_dir, stem + aux_ext))
            except FileNotFoundError:
                pass


def _flatten_sections(sections: dict) -> list[str]:
    """Flatten all section content into a single ordered list of strings."""
    items: list[str] = []
    for content in sections.values():
        if isinstance(content, str):
            items.extend(l for l in content.split("\n") if l.strip())
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, str):
                    items.append(item.strip())
                elif isinstance(item, dict):
                    for v in item.values():
                        if isinstance(v, str):
                            items.append(v.strip())
    return items


def _latex_escape(text: str) -> str:
    """Escape plain text for safe insertion into LaTeX source."""
    text = text.replace("\\", "\\textbackslash{}")
    text = text.replace("&", "\\&")
    text = text.replace("%", "\\%")
    text = text.replace("$", "\\$")
    text = text.replace("#", "\\#")
    text = text.replace("_", "\\_")
    text = text.replace("{", "\\{")
    text = text.replace("}", "\\}")
    text = text.replace("~", "\\textasciitilde{}")
    text = text.replace("^", "\\textasciicircum{}")
    return text


def _tex_replace(tex_source: str, original: str, edited: str) -> str:
    """
    Find `original` verbatim in the LaTeX source and replace with
    LaTeX-escaped `edited`. Because we now parse .tex without pandoc,
    `original` is extracted literally from the source and will match directly.
    Only replaces the first occurrence to avoid clobbering repeated phrases.

    Skills rows are stored as "Label: Value" but live in the source as
    \\textbf{Label} & Value — handled specially to preserve that structure.
    """
    # Skills row — original is "Label: Value" but source has \textbf{Label} & Value
    if ": " in original and original not in tex_source:
        colon_idx = original.index(": ")
        label_esc = re.escape(original[:colon_idx])
        old_val_esc = re.escape(original[colon_idx + 2:])
        pattern = r"\\textbf\{" + label_esc + r"\}\s*&\s*" + old_val_esc
        if re.search(pattern, tex_source):
            if ": " in edited:
                split_at = edited.index(": ")
                new_label = edited[:split_at]
                new_value = _latex_escape(edited[split_at + 2:])
            else:
                new_label = original[:colon_idx]
                new_value = _latex_escape(edited)
            replacement = f"\\textbf{{{new_label}}} & {new_value}"
            return re.sub(pattern, replacement, tex_source, count=1)

    replacement = _latex_escape(edited)
    if original in tex_source:
        return tex_source.replace(original, replacement, 1)
    return tex_source


TEXBIN = "/usr/local/texlive/2026basic/bin/universal-darwin"


def _pdflatex_compile(tex_path: str, out_dir: str) -> str:
    """Compile a .tex file to PDF. Tries pdflatex, xelatex, then lualatex."""
    candidates = [
        os.path.join(TEXBIN, "pdflatex"),
        os.path.join(TEXBIN, "xelatex"),
        os.path.join(TEXBIN, "lualatex"),
        "pdflatex", "xelatex", "lualatex",
    ]
    last_error = ""
    for compiler in candidates:
        try:
            result = subprocess.run(
                [compiler, "-interaction=nonstopmode", "-output-directory", out_dir, tex_path],
                capture_output=True,
                text=True,
                timeout=90,
            )
            pdf_path = os.path.join(out_dir, Path(tex_path).stem + ".pdf")
            if os.path.exists(pdf_path):
                return pdf_path
            # Compiler ran but produced no PDF — surface the error
            last_error = f"[{compiler}] exit {result.returncode}\n{result.stdout[-2000:]}\n{result.stderr[-1000:]}"
            break  # Don't try other compilers if this one was found but failed
        except FileNotFoundError:
            continue
        except subprocess.TimeoutExpired:
            last_error = f"[{compiler}] timed out"
            break

    raise RuntimeError(
        f"LaTeX compilation failed.\n{last_error}\n\n"
        "If compiler not found: brew install --cask basictex"
    )


# ---------------------------------------------------------------------------
# DOCX template export
# ---------------------------------------------------------------------------

def _export_docx_template(sections: dict, original_sections: dict | None, template_content: str, output_path: str) -> str:
    from docx import Document

    b64 = template_content.split(",", 1)[1] if "," in template_content else template_content
    template_bytes = base64.b64decode(b64)
    out_dir = os.path.dirname(output_path)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir="/tmp") as tmp:
        tmp.write(template_bytes)
        tmp_path = tmp.name

    filled_path = tmp_path.replace(".docx", "_filled.docx")
    try:
        doc = Document(tmp_path)
        if original_sections:
            _fill_docx_surgical(doc, original_sections, sections)
        else:
            _fill_docx_sections(doc, sections)
        doc.save(filled_path)
        pdf_path = _libreoffice_convert(filled_path, out_dir)
        os.rename(pdf_path, output_path)
        return output_path
    finally:
        for p in [tmp_path, filled_path]:
            try:
                os.unlink(p)
            except FileNotFoundError:
                pass


def _iter_all_paragraphs(doc):
    """
    Yield every paragraph in the document in reading order,
    including paragraphs inside table cells (which doc.paragraphs misses).
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table, _Cell

    def _walk(parent):
        if hasattr(parent, "element"):
            parent_elm = parent.element.body
        else:
            parent_elm = parent._tc  # _Cell

        for child in parent_elm.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, parent)
            elif child.tag == qn("w:tbl"):
                tbl = Table(child, parent)
                for row in tbl.rows:
                    for cell in row.cells:
                        yield from _walk(cell)

    yield from _walk(doc)


def _normalize_key(text: str) -> str:
    key = text.lower().strip()
    key = re.sub(r"[^a-z0-9\s]", "", key)
    key = re.sub(r"\s+", "_", key)
    return key


def _is_heading_para(para) -> bool:
    style_name = para.style.name.lower() if para.style else ""
    text = para.text.strip()
    if not text or len(text) >= 60:
        return False
    return "heading" in style_name or (
        para.runs
        and any(r.bold for r in para.runs)
        and (text.isupper() or text.istitle())
    )


def _strip_bullet_prefix(text: str) -> str:
    t = text.strip()
    for prefix in ("- ", "• ", "* "):
        if t.startswith(prefix):
            return t[len(prefix):]
    return t


def _clear_para_text(para) -> None:
    """Erase all text from a paragraph, even if runs aren't wrapped by python-docx."""
    if para.runs:
        para.runs[0].text = ""
        for run in para.runs[1:]:
            run.text = ""
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for t_elem in para._p.iter(f"{{{WNS}}}t"):
        t_elem.text = ""


def _fill_docx_surgical(doc, original_sections: dict, sections: dict):
    """
    Surgical text replacement: build an original→edited mapping from the two
    section dicts, then walk every paragraph in the template and swap matching
    text in-place.  Paragraph styles, indentation, spacing, and run-level
    formatting are preserved — only the text content changes.
    """
    orig_items = _flatten_sections(original_sections)
    edit_items = _flatten_sections(sections)

    # Build lookup: original_text → edited_text (first occurrence wins)
    replacements: dict[str, str] = {}
    for orig, edited in zip(orig_items, edit_items):
        if orig and orig not in replacements:
            replacements[orig] = edited

    if not replacements:
        return

    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for para in _iter_all_paragraphs(doc):
        text = para.text.strip()
        if not text or text not in replacements:
            continue
        new_text = _strip_bullet_prefix(replacements[text])
        # Clear ALL w:t elements in the paragraph XML tree first (handles runs
        # inside hyperlinks, tracked changes, and other nested structures).
        _clear_para_text(para)
        # Write new text into the first accessible location.
        if para.runs:
            para.runs[0].text = new_text
        else:
            # Text was inside a hyperlink or other nested element — find the
            # first w:t in the XML and write there directly.
            t_elems = list(para._p.iter(f"{{{WNS}}}t"))
            if t_elems:
                t_elems[0].text = new_text


def _fill_docx_sections(doc, sections: dict):
    """
    Sequential fill: iterate every paragraph in the template (including table cells),
    detect section headings, and replace content paragraph-by-paragraph.
    Leftover template paragraphs beyond the edited content are cleared.
    """
    content_lines: dict[str, list[str]] = {}
    for key, content in sections.items():
        lines: list[str] = []
        if isinstance(content, str):
            lines = [l for l in content.split("\n") if l.strip()]
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, str):
                    lines.append(item.strip())
                elif isinstance(item, dict):
                    for v in item.values():
                        if isinstance(v, str):
                            lines.append(v.strip())
        content_lines[key] = lines

    section_cursors: dict[str, int] = {k: 0 for k in content_lines}
    current_section: str | None = None

    for para in _iter_all_paragraphs(doc):
        text = para.text.strip()
        if not text:
            continue

        if _is_heading_para(para):
            current_section = _normalize_key(text)
            continue

        if not current_section or current_section not in content_lines:
            continue

        lines = content_lines[current_section]
        idx = section_cursors[current_section]

        if idx < len(lines):
            new_text = _strip_bullet_prefix(lines[idx])
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""
            section_cursors[current_section] = idx + 1
        else:
            _clear_para_text(para)


# ---------------------------------------------------------------------------
# Plain-docx fallback (no template)
# ---------------------------------------------------------------------------

def _export_plain_docx(sections: dict, output_path: str) -> str:
    from docx import Document

    doc = Document()
    for key, content in sections.items():
        display = key.replace("_", " ").title()
        doc.add_heading(display, level=1)

        if isinstance(content, str):
            doc.add_paragraph(content)
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, str):
                    if item.strip().startswith(("- ", "• ", "* ")):
                        p = doc.add_paragraph(style="List Bullet")
                        p.add_run(_strip_bullet_prefix(item))
                    else:
                        p = doc.add_paragraph()
                        p.add_run(item).bold = True
                elif isinstance(item, dict):
                    for k, v in item.items():
                        p = doc.add_paragraph()
                        p.add_run(f"{k.replace('_', ' ').title()}: ").bold = True
                        p.add_run(str(v))
        elif isinstance(content, dict):
            for k, v in content.items():
                p = doc.add_paragraph()
                p.add_run(f"{k.replace('_', ' ').title()}: ").bold = True
                p.add_run(str(v))

    out_dir = os.path.dirname(output_path)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir="/tmp") as tmp:
        tmp_path = tmp.name
    try:
        doc.save(tmp_path)
        pdf_path = _libreoffice_convert(tmp_path, out_dir)
        os.rename(pdf_path, output_path)
        return output_path
    finally:
        try:
            os.unlink(tmp_path)
        except FileNotFoundError:
            pass


# ---------------------------------------------------------------------------
# LibreOffice
# ---------------------------------------------------------------------------

def _libreoffice_convert(docx_path: str, out_dir: str) -> str:
    lo_cmd = LIBREOFFICE_BIN if os.path.exists(LIBREOFFICE_BIN) else "libreoffice"
    result = subprocess.run(
        [lo_cmd, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        capture_output=True,
        text=True,
        timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice conversion failed (exit {result.returncode}):\n"
            f"stdout: {result.stdout}\nstderr: {result.stderr}"
        )
    pdf_path = os.path.join(out_dir, Path(docx_path).stem + ".pdf")
    if not os.path.exists(pdf_path):
        raise RuntimeError(
            f"LibreOffice ran but PDF not found at {pdf_path}.\n"
            f"stdout: {result.stdout}\nstderr: {result.stderr}"
        )
    return pdf_path
